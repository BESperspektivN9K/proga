import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import PatternFill
import os
import subprocess
import math
import re
import openpyxl

def first_word_num(word_range):
    return int(word_range.split('-')[0]) if '-' in word_range else int(word_range)

def process_file(input_path, output_path):

    doc = Document(input_path)
    KTGS = 1.0
    wb = openpyxl.Workbook()
    ws_t0 = wb.active
    ws_t0.title = "T0"
    ws_20t0 = wb.create_sheet(title="20T0")
    ws_errors = wb.create_sheet(title="Праметры с ошибкой")
    ws_raz = wb.create_sheet(title="Разовые")
    ws_errors_raz = wb.create_sheet(title="Разовые с ошибкой")

    headers1 = ["Наименование", "Идентификатор", "Номер 16р слова", "Начальный участок выдачи", "Конечный участок выдачи", "Тип", "Номер разряда", "Задача/Блок", "Количество слов"]
    headers2 = ["Наименование", "Идентификатор", "Номер 16р слова", "Номер счетчика", "Начальный участок выдачи", "Конечный участок выдачи", "Тип", "Задача/Блок", "Количество слов"]
    headers3 = ["Наименование", "Идентификатор", "Номер 16р слова", "Номер участка", "Номер счетчика", "Тип", "Номер разряда", "Задача/Блок", "Количество слов"]
    ws_t0.append(headers1)
    ws_20t0.append(headers2)
    ws_errors.append(headers1)
    ws_raz.append(headers3)
    ws_errors_raz.append(headers3)

    words = []
    used_words = []
    rows_T0 = []#список строк Т0
    rows_20T0 = []#список строк 20Т0
    rows_T0_raz = []#список разрядных строк Т0
    rows_razovie = []#список строк разовых
    rows_razrad_raz = []#список строк разовых разрядных
    excel_rowT0 = 1
    excel_row_razT0 = 1
    excel_row_raz = 1
    excel_row_Razrad = 0
    rows_errors = []# параметры Т0 и 20Т0 с ошибкой
    rows_errors_raz = []# параметры разовые с ошибкой
    available_time = [100,101,105,110,120,125,130,135,140,145,150,155,160,165,170,175,180,185,190,195,200,205,210,215,220,225,230,240,255]
    available_tip = [1,2,3,4,5,6,7,8]
    available_uchastki = []
    number_word = 51

    for i, row in enumerate(doc.tables[0].rows):
        if i == 0:
            continue

        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 5:
            continue

        name = cells[0]
        id = cells[1]
        group = cells[2]

        numbers = re.findall(r"\d+", cells[3])
        if len(numbers) == 3:
            start_num = numbers[1]
            end_num = numbers[2]
        else:
            start_num = numbers[0]
            end_num = numbers[1]

        words_to_add = 1.0
        if 'р' in cells[-2]:
            match = re.findall(r'(\d+)р', cells[-2])
            kolvo = float(match[-1])
            words_to_add = float(kolvo) / 16.0
        elif ',' in cells[-2]:
            normal_value = cells[-2].replace(',', '.')
            kolvo = normal_value.split('/')[0]
            numbers = re.findall(r'\d+\.?\d*', kolvo)
            for num in numbers:
                words_to_add *= float(num)
        else:
            kolvo = cells[-2].split('/')[0]
            numbers = re.findall(r'(\d+)', kolvo)
            for num in numbers:
                words_to_add *= float(num)

        if 'Ктгс' in cells[-2]:
            words_to_add *= KTGS

        tip_match = re.findall(r"/(\d+)", cells[-2])
        tip = tip_match[0] if tip_match else "8"
        tip = tip[0] if tip else "8"

        block = cells[-1]

        if group == "T0":
            if (isinstance(words_to_add, float) and words_to_add != int(words_to_add)):
                excel_row_razT0 += 1
                excel_row = excel_row_razT0
            else:
                target_ws = ws_t0
                excel_rowT0 += 1
                excel_row = excel_rowT0
        else:
            target_ws = ws_20t0
            excel_row = target_ws.max_row + 1

        row_data = {
            "name": name,
            "id": id,
            "group": group,
            "start_num": start_num,
            "end_num": end_num,
            "tip": tip,
            "block": block,
            "words_to_add": words_to_add,
            "excel_row": excel_row
        }

        if int(start_num) not in available_time or int(end_num) not in available_time or int(tip) not in available_tip or group not in ['T0', '20T0'] or (int(start_num) >= int(end_num)):
            rows_errors.append(row_data)
        else:
            if group == "20T0":
                rows_20T0.append(row_data)
                target_ws.append([name, id, "", "", start_num, end_num, tip, block, words_to_add])
            else:
                if (isinstance(words_to_add, float) and words_to_add != int(words_to_add)):
                    rows_T0_raz.append(row_data)
                else:
                    rows_T0.append(row_data)

    for i, row in enumerate(doc.tables[1].rows):
        if i == 0:
            continue

        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 5:
            continue

        name = cells[0]
        id = cells[1]
    
        numbers = re.findall(r"\d+", cells[2])
        uchastok = numbers[-1]

        words_to_add = 1.0
        if 'р' in cells[-2]:
            match = re.findall(r'(\d+)р', cells[-2])
            kolvo = float(match[-1])
            words_to_add = float(kolvo) / 16.0
        elif ',' in cells[-2]:
            normal_value = cells[-2].replace(',', '.')
            kolvo = normal_value.split('/')[0]
            numbers = re.findall(r'\d+\.?\d*', kolvo)
            for num in numbers:
                words_to_add *= float(num)
        else:
            kolvo = cells[-2].split('/')[0]
            numbers = re.findall(r'(\d+)', kolvo)
            for num in numbers:
                words_to_add *= float(num)

        if 'Ктгс' in cells[-2]:
            words_to_add *= KTGS

        tip_match = re.findall(r"/(\d+)", cells[-2])
        tip = tip_match[0] if tip_match else "8"
        tip = tip[0] if tip else "8"

        block = cells[-1]

        if (isinstance(words_to_add, float) and words_to_add != int(words_to_add)):
            excel_row_raz +=1
            excel_row = excel_row_raz
        else:
            excel_row_Razrad += 1
            excel_row = excel_row_Razrad

        row_data = {
            "name": name,
            "id": id,
            "uchastok": uchastok,
            "tip": tip,
            "block": block,
            "words_to_add": words_to_add,
            "excel_row": excel_row
        }

        if int(tip) not in available_tip:#добавить сюда участки
            rows_errors_raz.append(row_data)
        else:
            if (isinstance(words_to_add, float) and words_to_add != int(words_to_add)):
                rows_razrad_raz.append(row_data)
            else:
                 rows_razovie.append(row_data)

# ---- Сначала распределяем слова для 20T0 ----

    total_words_20t0 = sum(int(row["words_to_add"]) for row in rows_20T0)
    count_words = 2 if total_words_20t0 < 20 else math.ceil(total_words_20t0 / 20)

    counter = 1
    max_counters = 20
    count_words_per_counter = count_words
    available_words = [str(number_word + i) for i in range(count_words_per_counter)]
    used_words_per_counter = {i: [] for i in range(1, max_counters + 1)}
    final_rows_20T0 = []

    for row in rows_20T0:
        name = row["name"]
        id = row["id"]
        start_num = int(row["start_num"])
        end_num = int(row["end_num"])
        tip = row["tip"]
        block = row["block"]
        words_needed = int(row["words_to_add"])
        counter = 1

        while words_needed > 0 and counter <= max_counters:
            used_in_counter = used_words_per_counter[counter]

            free_words = []
            for word in available_words:
                conflict = False
                for used in used_in_counter:
                    if word == used["номер"]:
                        if not (end_num <= used["start"] or start_num >= used["end"]):
                            conflict = True
                            break
                if not conflict:
                    free_words.append(word)

            if not free_words:
                counter += 1
                continue

            take = min(words_needed, len(free_words))
            taken_words = free_words[:take]

            for word in taken_words:
                used_words_per_counter[counter].append({
                    "номер": word,
                    "start": start_num,
                    "end": end_num
                })

            word_range = taken_words[0] if len(taken_words) == 1 else f"{taken_words[0]}-{taken_words[-1]}"

            final_rows_20T0.append({
                "name": name,
                "id": id,
                "word_range": word_range,
                "counter": counter,
                "start_num": start_num,
                "end_num": end_num,
                "tip": tip,
                "block": block,
                "count": take
            })

            words_needed -= take
            counter += 1

        if words_needed > 0:
            raise RuntimeError(f"Не удалось назначить все слова для параметра: {name}")

    ws_20t0.delete_rows(2, ws_20t0.max_row)
    final_rows_20T0.sort(key=lambda x: (x["counter"], first_word_num(x["word_range"])))

    for row in final_rows_20T0:
        ws_20t0.append([
            row["name"],
            row["id"],
            row["word_range"],
            row["counter"],
            row["start_num"],
            row["end_num"],
            row["tip"],
            row["block"],
            row["count"]
        ])

    number_word += count_words
    max_counter = max(row["counter"] for row in final_rows_20T0)
    ws_20t0.merge_cells('K3:S3')
    ws_20t0.merge_cells('K4:S4')
    ws_20t0['K3'] = f"Количество слов, выделенное под параметры 20Т0: {count_words} ({number_word-count_words} - {number_word-1})"
    ws_20t0['K4'] = f"Количество счетчиков, использованных для распределения параметров 20Т0:  {max_counter}"

    def min_proper_divisor(n):
        for i in range(2, n + 1):
            if n % i == 0:
                return i
        return n  # если n — простое число

    total_words_razovie = sum(int(row["words_to_add"]) for row in rows_razovie)
    count_words_razovie = min_proper_divisor(total_words_razovie)

    available_words_razovie = [str(number_word + i) for i in range(count_words_razovie)]
    used_words_per_counter_razovie = {i: {} for i in range(1, 100)}  # счетчиков может быть сколько угодно

    final_rows_razovie = []

    counter = 1
    for row in rows_razovie:
        name = row["name"]
        id = row["id"]
        uchastok = row["uchastok"]
        tip = row["tip"]
        block = row["block"]
        words_needed = int(row["words_to_add"])

        counter = 1
        while words_needed > 0:
            used_in_counter = used_words_per_counter_razovie[counter]

            free_words = []
            for word in available_words_razovie:
                used_uchastki = used_in_counter.get(word, set())
                if uchastok not in used_uchastki:
                    free_words.append(word)

            if not free_words:
                counter += 1
                continue

            take = min(words_needed, len(free_words))
            taken_words = free_words[:take]

            for word in taken_words:
                used_words_per_counter_razovie[counter].setdefault(word, set()).add(uchastok)

            word_range = taken_words[0] if len(taken_words) == 1 else f"{taken_words[0]}-{taken_words[-1]}"

            final_rows_razovie.append({
                "name": name,
                "id": id,
                "word_range": word_range,
                "counter": counter,
                "uchastok": uchastok,
                "tip": tip,
                "block": block,
                "count": take
            })

            words_needed -= take
            counter += 1

        if words_needed > 0:
            raise RuntimeError(f"Не удалось назначить все слова для разового параметра: {name}")

# Очистка листа и запись
        ws_raz.delete_rows(2, ws_raz.max_row)
        final_rows_razovie.sort(key=lambda x: (x["counter"], x["uchastok"]))
        
        for row in final_rows_razovie:
            ws_raz.append([
                row["name"],
                row["id"],
                row["word_range"],
                row["uchastok"],
                row["counter"],
                row["tip"],
                "",
                row["block"],
                row["count"]
            ])
    
    rows_razrad_raz.sort(key=lambda row: row["uchastok"])
    for row in rows_razrad_raz:
        ws_raz.append([
            row["name"],
            row["id"],
            "",
            row["uchastok"],
            "",
            "",  # ...
            row["tip"],
            row["block"],
            row["words_to_add"]
        ])
        row["excel_row"] = ws_raz.max_row
        
    b, k = 0, 1
    counter = max(row["counter"] for row in final_rows_razovie) + 1 
    while b < len(rows_razrad_raz):
        row = rows_razrad_raz[b]
        razrad = row["words_to_add"] * 16
        if k + razrad > 17:
            counter += 1
            k = 1
        else:
            ws_raz.cell(row=row["excel_row"], column=3).value = number_word
            ws_raz.cell(row=row["excel_row"], column=5).value = counter
            ws_raz.cell(row=row["excel_row"], column=7).value = f"{int(k)}р" if razrad == 1 else f"{int(k)}-{int(k + razrad - 1)}р"
            k += razrad
            b += 1
    number_word += count_words_razovie
    ws_raz.merge_cells('K3:S3')
    ws_raz.merge_cells('K4:S4')
    ws_raz['K3'] = f"Количество слов, выделенное под разовые параметры: {count_words_razovie} ({number_word-count_words_razovie} - {number_word-1})"
    ws_raz['K4'] = f"Количество счетчиков, использованных для распределения разовых параметров:  {counter}"

    ################################################### ----  выдача слов для T0 ----
    for row in rows_T0:
        ws_t0.append([
            row["name"],
            row["id"],
            "",
            row["start_num"],
            row["end_num"],
            row["tip"],
            "",
            row["block"],
            row["words_to_add"]
        ])
        row["excel_row"] = ws_t0.max_row

    for row in rows_T0:
        start_num = row["start_num"]
        end_num = row["end_num"]
        words_to_add = row["words_to_add"]
        excel_row = row["excel_row"]
        first_word = number_word

        for _ in range(int(words_to_add)):
            words.append({
                "номер": str(number_word),
                "начальный": start_num,
                "конечный": end_num
            })
            used_words.append({
                "номер": str(number_word),
                "начальный": start_num,
                "конечный": end_num
            })
            number_word += 1

        last_word = number_word - 1
        word_range = f"{first_word}" if first_word == last_word else f"{first_word}-{last_word}"
        ws_t0.cell(row=excel_row, column=3).value = word_range
########################################################
    
    rows_T0_raz.sort(key=lambda row: row["start_num"])
    for row in rows_T0_raz:
        ws_t0.append([
            row["name"],
            row["id"],
            "",
            row["start_num"],
            row["end_num"],
            row["tip"] if row["tip"] == '7' else "",
            "",
            row["block"],
            row["words_to_add"]
        ])
        row["excel_row"] = ws_t0.max_row
        
    b, k = 0, 1
    
    while b < len(rows_T0_raz):
        row = rows_T0_raz[b]
        razrad = row["words_to_add"] * 16
        if k + razrad > 17:
            number_word += 1
            k = 1
        else:
            ws_t0.cell(row=row["excel_row"], column=3).value = number_word
            ws_t0.cell(row=row["excel_row"], column=7).value = f"{int(k)}р" if razrad == 1 else f"{int(k)}-{int(k + razrad - 1)}р"
            k += razrad
            b += 1

    red_fill = PatternFill(start_color="FF9999", end_color="FF9999", fill_type="solid")
    for row in rows_errors:
        ws_errors.append([
            row["name"],
            row["id"],
            "",
            row["start_num"],
            row["end_num"],
            row["tip"],
            row["block"],
            row["words_to_add"]
        ])
        last_row = ws_errors.max_row
        for col in range(1, 9):
            ws_errors.cell(row=last_row, column=col).fill = red_fill
    
#######################################################################
    wb.save(output_path)

def select_input_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    input_entry.delete(0, tk.END)
    input_entry.insert(0, file_path)


def select_output_file():
    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
    output_entry.delete(0, tk.END)
    output_entry.insert(0, file_path)


def run_program():
    input_path = input_entry.get()
    output_path = output_entry.get()

    if not input_path:
        messagebox.showwarning("Предупреждение", "Пожалуйста, выберите входной файл.")
        return
    if not output_path:
        messagebox.showwarning("Предупреждение", "Пожалуйста, выберите путь для сохранения файла.")
        return

    try:
        process_file(input_path, output_path)
        messagebox.showinfo("Готово", "Файл успешно обработан и сохранён.")
        os.startfile(output_path)  # Открыть файл после завершения
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка:\n{str(e)}")

################################################################### --- Интерфейс ---
root = tk.Tk()
root.title("Обработка параметров")

tk.Label(root, text="Входной Word файл:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
input_entry = tk.Entry(root, width=50)
input_entry.grid(row=0, column=1, padx=5, pady=5)
tk.Button(root, text="Обзор", command=lambda: input_entry.insert(0, filedialog.askopenfilename(filetypes=[("Word files", "*.docx")]))).grid(row=0, column=2, padx=5, pady=5)

tk.Label(root, text="Сохранить Excel как:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
output_entry = tk.Entry(root, width=50)
output_entry.grid(row=1, column=1, padx=5, pady=5)
tk.Button(root, text="Обзор", command=lambda: output_entry.insert(0, filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")]))).grid(row=1, column=2, padx=5, pady=5)

tk.Button(root, text="Запустить", command=lambda: run_program(), width=20, bg="green", fg="white").grid(row=2, column=1, pady=15)

root.mainloop()
